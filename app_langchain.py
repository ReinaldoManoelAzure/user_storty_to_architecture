import os
import re
import io
import base64
import asyncio
import nest_asyncio
from datetime import datetime

import streamlit as st
import requests
import docx
from PyPDF2 import PdfReader
from dotenv import load_dotenv

# LangChain + Gemini
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter

# =============================
# Correções de ambiente para Streamlit + gRPC async
# =============================
nest_asyncio.apply()
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

# =============================
# Carregar variáveis de ambiente
# =============================
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# =============================
# Utilidades: Extração de texto
# =============================
def extract_text_from_file(uploaded_file):
    if uploaded_file.name.lower().endswith(".docx"):
        d = docx.Document(uploaded_file)
        return "\n".join([p.text for p in d.paragraphs])
    elif uploaded_file.name.lower().endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text.append(t)
        return "\n".join(text)
    else:
        st.warning("Formato de arquivo não suportado. Use PDF ou DOCX.")
        return None

# =============================
# PlantUML Helpers
# =============================
def encode_6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    if b == 0: return '-'
    if b == 1: return '_'
    return '?'

def append3bytes(b1, b2, b3):
    c1 = b1 >> 2
    c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
    c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
    c4 = b3 & 0x3F
    return "".join([
        encode_6bit(c1 & 0x3F),
        encode_6bit(c2 & 0x3F),
        encode_6bit(c3 & 0x3F),
        encode_6bit(c4 & 0x3F),
    ])

def encode_plantuml(text):
    import zlib
    data = zlib.compress(text.encode("utf-8"))[2:-4]
    res, i = "", 0
    while i < len(data):
        b1 = data[i]
        b2 = data[i + 1] if i + 1 < len(data) else 0
        b3 = data[i + 2] if i + 2 < len(data) else 0
        res += append3bytes(b1, b2, b3)
        i += 3
    return res

def extract_plantuml_code(text: str) -> str | None:
    match = re.search(r'@startuml.*?@enduml', text, re.DOTALL | re.IGNORECASE)
    return match.group(0).strip() if match else None

def plantuml_get_diagram(plantuml_code, fmt="svg"):
    server_url = f"http://www.plantuml.com/plantuml/{fmt}/"
    encoded = encode_plantuml(plantuml_code)
    url = server_url + encoded
    response = requests.get(url)
    if response.status_code == 200:
        return response.content if fmt == "png" else response.text
    else:
        st.warning(f"Não foi possível gerar o diagrama PlantUML em {fmt}.")
        return None

# =============================
# Sanitização de PlantUML (C4 estrito)
# =============================
FORBIDDEN_TOKENS = [
    r'^\s*rectangle\b', r'^\s*note\b', r'^\s*legend\b', r'^\s*class\b', r'^\s*interface\b',
    r'^\s*entity\b', r'^\s*database\b', r'^\s*component\b(?!\s*\()'
]
ALLOWED_C4_PREFIXES = [
    'Person(', 'Person_Ext(', 'System_Boundary(', 'System(', 'System_Ext(',
    'Container(', 'ContainerDb(', 'ContainerQueue(', 'Container_Ext(', 'ContainerDb_Ext(', 'ContainerQueue_Ext(',
    'Component(', 'ComponentDb(', 'ComponentQueue(', 'Component_Ext(', 'ComponentDb_Ext(', 'ComponentQueue_Ext(',
    'Rel(', 'Rel_U(', 'Rel_D(', 'Rel_L(', 'Rel_R('
]

def sanitize_plantuml_c4(uml: str):
    if not uml:
        return uml, False, []

    adjusted = False
    removed_lines = []

    lines = uml.splitlines()
    clean = []

    for ln in lines:
        if re.match(r'^\s*@startuml', ln) or re.match(r'^\s*@enduml', ln) or re.match(r'^\s*!include', ln):
            clean.append(ln)
            continue
        if any(re.search(tok, ln, flags=re.IGNORECASE) for tok in FORBIDDEN_TOKENS):
            adjusted = True
            removed_lines.append(ln)
            continue
        if ln.strip() == "" or ln.strip().startswith("'"):
            clean.append(ln)
            continue
        if any(ln.strip().startswith(prefix) for prefix in ALLOWED_C4_PREFIXES):
            clean.append(ln)
            continue
        if re.match(r'^\s*skinparam\b', ln, flags=re.IGNORECASE):
            clean.append(ln)
            continue

        adjusted = True
        removed_lines.append(ln)

    content_lines = [l for l in clean if not re.match(r'^\s*@startuml|^\s*@enduml', l)]
    if len(content_lines) == 0 and len(clean) >= 2:
        return uml, False, []

    return "\n".join(clean), adjusted, removed_lines

# =============================
# Normalização do Markdown
# =============================
def normalize_markdown(md: str) -> str:
    if not md:
        return md

    md = re.sub(r'\[\]\(.*?\)', '', md)
    md = re.sub(r'(\*\*\d\)\s*[^\*]+?\*\*)([ \t]*)(?!\n\n)', r'\1\n\n', md)
    md = re.sub(r'(Data:\s*[^\n]+)\s*', r'\1\n', md)
    md = re.sub(r'(Autor:\s*[^\n]+)\s*', r'\1\n', md)
    md = re.sub(r'(Versão:\s*[^\n]+)\s*', r'\1\n', md)

    return md.strip()

# =============================
# Embeddings + Mini-RAG
# =============================
def make_embeddings():
    return GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=GOOGLE_API_KEY,
    )

def build_index(text: str):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    parts = splitter.split_text(text)
    docs = [Document(page_content=p) for p in parts]
    embeddings = make_embeddings()
    return FAISS.from_documents(docs, embeddings)

def retrieve_relevant_chunks(index, query, k=4):
    try:
        return index.similarity_search(query, k=k)
    except Exception:
        return []

# =============================
# LLM - Gemini com fallback
# =============================
PRIMARY_MODEL = "gemini-2.5-flash"
FALLBACK_MODEL = "gemini-1.5-flash"

def _mk_llm(model_name, temp=0.25, max_tokens=6000):
    return ChatGoogleGenerativeAI(
        model=model_name,
        temperature=temp,
        max_output_tokens=max_tokens,
        google_api_key=GOOGLE_API_KEY,
    )

def gemini_with_fallback():
    try:
        return _mk_llm(PRIMARY_MODEL)
    except Exception:
        return _mk_llm(FALLBACK_MODEL)

# =============================
# Prompt de arquitetura padronizado
# =============================
def make_arch_prompt(text: str, rules_context: str) -> str:
    today = datetime.today().strftime("%d/%m/%Y")
    return f"""
Você é um **Arquiteto(a) de Software e Soluções** responsável por elaborar **documentação de arquitetura técnica**.  
Sua missão é transformar o texto abaixo em um documento **padronizado, claro e completo**, que apoie **desenvolvedores, analistas e stakeholders** a compreenderem a solução proposta sob a ótica arquitetural.  

⚠️ **REGRAS OBRIGATÓRIAS DE FORMATAÇÃO**  
- O título do documento deve ser exatamente:  
  ## **Documento de Arquitetura de Software e Soluções – [Nome do Sistema]**  
- Abaixo do título insira uma linha separadora: ---  
- Logo depois, insira exatamente estas 3 linhas (sem hyperlinks nem markdown adicional):  
  Data: {today}  
  Autor: Arquiteto(a) de Software e Soluções  
  Versão: 1.0  

---

**1) Objetivo**  

**2) Contexto e Requisitos**  

**3) Arquitetura e Diagrama de Componentes (C4-PlantUML)**  

**4) Integrações Externas e Contratos em Alto Nível**  

**5) Decisões Arquiteturais Relevantes**  

**6) Critérios de Aceite e Próximos Passos**  

---

=== REGRAS CORPORATIVAS RELEVANTES (contexto) ===
{rules_context}

=== TEXTO BASE ===
{text}
""".strip()

# =============================
# Geração do documento
# =============================
def generate_full_doc(user_text: str, corporate_rules: list[str] | None = None, progress_cb=None) -> str:
    if progress_cb: progress_cb(10, "Criando índice de regras (mini-RAG)...")

    relevant_rules_text = ""
    if corporate_rules:
        rules_text = "\n".join([r.strip() for r in corporate_rules if r.strip()])
        if rules_text:
            try:
                idx = build_index(rules_text)
                chunks = retrieve_relevant_chunks(idx, user_text, k=6)
                relevant_rules_text = "\n".join([c.page_content for c in chunks]) if chunks else rules_text
            except Exception:
                relevant_rules_text = rules_text

    if progress_cb: progress_cb(35, "Montando prompt...")
    prompt = make_arch_prompt(user_text, relevant_rules_text)

    if progress_cb: progress_cb(55, "Chamando LLM (Gemini) com fallback...")
    llm = gemini_with_fallback()
    try:
        resp = llm.invoke(prompt)
        raw_md = resp.content if hasattr(resp, "content") else str(resp)
    except Exception as e:
        raise RuntimeError(f"Falha na chamada ao LLM: {e}")

    if progress_cb: progress_cb(75, "Normalizando Markdown...")
    md = normalize_markdown(raw_md)

    if progress_cb: progress_cb(85, "Verificando e sanitizando diagrama C4 (se presente)...")
    plantuml = extract_plantuml_code(md)
    sanitized_info = {"adjusted": False, "removed": []}
    if plantuml:
        clean_uml, adjusted, removed = sanitize_plantuml_c4(plantuml)
        if adjusted:
            md = md.replace(plantuml, clean_uml)
            sanitized_info["adjusted"] = True
            sanitized_info["removed"] = removed

    if progress_cb: progress_cb(100, "Concluído.")
    return md, sanitized_info

# =============================
# UI Streamlit
# =============================
st.set_page_config(page_title="Gerador de Documento Técnico", layout="wide")
st.title("📄 Gerador de Documento Técnico Consolidado")

uploaded_file = st.file_uploader("📥 Envie um arquivo PDF ou DOCX (opcional):", type=["pdf", "docx"])
texto_digitado = st.text_area("✍️ Ou cole aqui a história de usuário/caso de uso:", height=220)

with st.expander("🏛️ Regras Corporativas (opcional) — usadas na análise de aderência", expanded=True):
    default_rules = (
        "Todas as integrações devem passar pelo API Gateway corporativo (ex.: Sensedia) com OAuth2.\n"
        "Logs de auditoria devem ser centralizados e imutáveis por 5 anos.\n"
        "Dados sensíveis devem estar cifrados em trânsito (TLS 1.2+) e em repouso (KMS corporativo).\n"
        "Serviços devem expor métricas (Prometheus/OpenTelemetry) e tracing distribuído.\n"
        "Preferência por comunicação assíncrona para integrações de longa duração (fila/stream).\n"
        "Endpoints públicos devem ter WAF e rate limiting.\n"
        "Backups diários com teste de restauração semanal.\n"
        "Segregar ambientes (dev/hml/prd) com contas/projetos distintos e IaC versionado.\n"
    )
    rules_text = st.text_area("Edite as regras (uma por linha):", value=default_rules, height=160)
    corporate_rules = [r.strip() for r in rules_text.splitlines() if r.strip()]

col1, col2 = st.columns([1, 1])
with col1:
    run = st.button("🚀 Gerar Documento")
with col2:
    reset = st.button("🧼 Limpar")

if reset:
    st.session_state.pop("documento_md", None)
    st.session_state.pop("sanitized_info", None)
    st.rerun()

if run:
    texto_base = None
    if uploaded_file:
        texto_base = extract_text_from_file(uploaded_file)
    elif texto_digitado and texto_digitado.strip():
        texto_base = texto_digitado.strip()

    if not GOOGLE_API_KEY:
        st.error("⚠️ GOOGLE_API_KEY não definido. Configure no .env ou variável de ambiente.")
    elif not texto_base:
        st.warning("Forneça um texto ou arquivo para gerar o documento.")
    else:
        progress = st.progress(0, text="Inicializando...")
        def set_progress(p, label):
            progress.progress(p, text=label)

        try:
            md, sanitized_info = generate_full_doc(texto_base, corporate_rules, progress_cb=set_progress)
            st.session_state["documento_md"] = md
            st.session_state["sanitized_info"] = sanitized_info
            st.success("Documento gerado com sucesso!")
        except Exception as e:
            st.error(f"Falha ao gerar documento: {e}")

# =============================
# Visualização
# =============================
if "documento_md" in st.session_state:
    st.subheader("📑 Documento Consolidado")
    st.markdown(st.session_state["documento_md"])

    plantuml_code = extract_plantuml_code(st.session_state["documento_md"])
    if plantuml_code:
        st.divider()
        st.markdown("### 🔎 Diagrama Arquitetural (C4-PlantUML)")

        svg_content = plantuml_get_diagram(plantuml_code, "svg")

        def render_svg_in_component(svg_text: str, height: int = 640):
            try:
                svg_clean = re.sub(r'<\?xml[^>]*\?>', '', svg_text, flags=re.IGNORECASE | re.DOTALL)
                svg_clean = re.sub(r'<!DOCTYPE[^>]*>', '', svg_clean, flags=re.IGNORECASE | re.DOTALL)
                svg_clean = svg_clean.strip()
                html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body style='margin:0'>{svg_clean}</body></html>"
                st.components.v1.html(html, height=height, scrolling=True)
                return True
            except Exception:
                return False

        rendered = False
        if svg_content:
            rendered = render_svg_in_component(svg_content, height=640)

