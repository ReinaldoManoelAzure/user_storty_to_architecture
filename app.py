import streamlit as st
import os
import re
import requests
import docx
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from llama_index.llms.google_genai import GoogleGenAI
import io
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.units import cm
import base64
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import VectorStoreIndex, Document, StorageContext
from llama_index.core.node_parser import SimpleNodeParser
from datetime import datetime


load_dotenv()

# --- Funções de download ---
def gerar_download_markdown(documento: str):
    st.download_button(
        label="⬇️ Baixar em Markdown",
        data=documento.encode("utf-8"),
        file_name="documento.md",
        mime="text/markdown"
    )

def gerar_download_pdf(documento: str):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # separa se há diagrama UML
    plantuml_code = extract_plantuml_code(documento)
    if plantuml_code:
        documento_sem_uml = documento.replace(plantuml_code, "[DIAGRAMA_ARQUITETURAL]")
    else:
        documento_sem_uml = documento

    for linha in documento_sem_uml.split("\n"):
        if linha.strip() == "[DIAGRAMA_ARQUITETURAL]":
            png = plantuml_get_diagram(plantuml_code, "png")
            if png:
                img_buffer = io.BytesIO(png)
                story.append(Image(img_buffer, width=14*cm, height=8*cm))
                story.append(Spacer(1, 12))
        else:
            if linha.strip() == "":
                story.append(Spacer(1, 12))
            else:
                story.append(Paragraph(linha, styles["Normal"]))

    doc.build(story)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Baixar em PDF",
        data=buffer,
        file_name="documento.pdf",
        mime="application/pdf"
    )


# =============================
# Extração de texto de arquivos
# =============================
def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif uploaded_file.name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    else:
        st.warning("Formato de arquivo não suportado. Use PDF ou DOCX.")
        return None

# =============================
# Encoder oficial PlantUML
# =============================
def encode_6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    if b == 0: return '-'
    if b == 1: return '_'
    return '?'

def append3bytes(b1, b2, b3):
    c1 = b1 >> 2
    c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
    c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
    c4 = b3 & 0x3F
    return "".join([
        encode_6bit(c1 & 0x3F),
        encode_6bit(c2 & 0x3F),
        encode_6bit(c3 & 0x3F),
        encode_6bit(c4 & 0x3F),
    ])

def encode_plantuml(text):
    import zlib
    data = zlib.compress(text.encode("utf-8"))[2:-4]
    res, i = "", 0
    while i < len(data):
        b1 = data[i]
        b2 = data[i + 1] if i + 1 < len(data) else 0
        b3 = data[i + 2] if i + 2 < len(data) else 0
        res += append3bytes(b1, b2, b3)
        i += 3
    return res

def extract_plantuml_code(text: str) -> str | None:
    match = re.search(r'@startuml.*?@enduml', text, re.DOTALL | re.IGNORECASE)
    return match.group(0).strip() if match else None

def plantuml_get_diagram(plantuml_code, fmt="svg"):
    server_url = f"http://www.plantuml.com/plantuml/{fmt}/"
    encoded = encode_plantuml(plantuml_code)
    url = server_url + encoded
    response = requests.get(url)
    if response.status_code == 200:
        return response.content if fmt == "png" else response.text
    else:
        st.warning(f"Não foi possível gerar o diagrama PlantUML em {fmt}.")
        return None

# =============================
# LLM - Gemini
# =============================
def _mk_llm(model="gemini-2.5-flash", temp=0.25, max_tokens=6000):
    return GoogleGenAI(
        model=model,
        temperature=temp,
        max_output_tokens=max_tokens,
    )

def safe_complete(prompt, temp=0.25, max_tokens=4000):
    """Executa chamada ao Gemini com fallback automático."""
    try:
        llm = _mk_llm("gemini-2.5-flash", temp=temp, max_tokens=max_tokens)
        return llm.complete(prompt).text
    except Exception as e:
        st.warning("⚠️ O modelo Gemini 2.5 está sobrecarregado. Tentando fallback para Gemini 1.5...")
        try:
            llm_fallback = _mk_llm("gemini-1.5-flash", temp=temp, max_tokens=max_tokens)
            return llm_fallback.complete(prompt).text
        except Exception as e2:
            st.error("❌ Não foi possível gerar a resposta. Tente novamente em alguns minutos.")
            return ""

# =============================
# Helpers para divisão de texto
# =============================
def split_text(text, max_chars=3000):
    for i in range(0, len(text), max_chars):
        yield text[i:i+max_chars]

# =============================
# >>> Mini-RAG para textos grandes
# =============================
def build_index(text: str):
    parser = SimpleNodeParser.from_defaults(chunk_size=512, chunk_overlap=50)
    docs = [Document(text)]
    nodes = parser.get_nodes_from_documents(docs)
    embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
    index = VectorStoreIndex(
        nodes,
        embed_model=embed_model,
        storage_context=StorageContext.from_defaults()
    )
    return index

def retrieve_relevant_chunks(index, query: str, top_k: int = 5) -> str:
    retriever = index.as_retriever(similarity_top_k=top_k)
    nodes = retriever.retrieve(query)
    return "\n".join([n.text for n in nodes])

# =============================
# Funções de geração parcial
# =============================
def generate_requisitos_doc(text):
    prompt = f"""
Você é um analista de requisitos.
Extraia APENAS os **Requisitos Funcionais, Não Funcionais e Regras de Negócio** do texto.
Responda de forma objetiva, sem introduções, conclusões ou comentários adicionais.
Organize em tópicos claros e numerados, diretamente.

=== TEXTO ===
{text}
"""
    return safe_complete(prompt, temp=0.25, max_tokens=2000)


def generate_technical_doc(text, stacks=""):
    
    today = datetime.today().strftime("%d/%m/%Y")

    prompt_content = f"""
Você é um **Arquiteto(a) de Software e Soluções** responsável por elaborar **documentação de arquitetura técnica**.  
Sua missão é transformar o texto abaixo em um documento **padronizado, claro e completo**, que apoie **desenvolvedores, analistas e stakeholders** a compreenderem a solução proposta sob a ótica arquitetural.  

⚠️ **REGRAS OBRIGATÓRIAS DE FORMATAÇÃO**  
- O título do documento deve ser exatamente:  
  ## **Documento de Arquitetura de Software e Soluções – [Nome do Sistema]**  
- Abaixo do título insira uma linha separadora: ---  
- Logo depois, insira exatamente estas 3 linhas (sem hyperlinks nem markdown adicional):  
  Data: {today}  
  Autor: Arquiteto(a) de Software e Soluções  
  Versão: 1.0  
- Sempre insira uma **linha em branco** após cada cabeçalho de seção.  
- ❌ Nunca escreva conteúdo na mesma linha do cabeçalho.  
- ❌ Nunca adicione frases extras (ex.: “Com certeza!”, “Aqui está...”) fora da estrutura.  
- ❌ Nunca altere a ordem ou os títulos das seções.  
- ✅ O conteúdo deve ser direto, objetivo e técnico.  

---

### 📑 Estrutura obrigatória do documento

**1) Objetivo**  

Descreva de forma clara o propósito do sistema, seu escopo e a visão arquitetural de alto nível.  

**2) Contexto e Requisitos**  

- **Contexto**: Explique o problema/necessidade que o sistema resolve.  
- **Requisitos Funcionais (RF)**: Liste numerada (RF001, RF002, …).  
- **Requisitos Não Funcionais (RNF)**:  
  - Liste numerada (RNF001, RNF002, …).  
  - Associe explicitamente cada RNF aos componentes relevantes do diagrama usando a sintaxe:  
    RNF00X – [Categoria] – Componentes afetados: [ComponenteA, ComponenteB] – Descrição objetiva.  

**3) Arquitetura e Diagrama de Componentes (C4-PlantUML)**  

Inclua um diagrama C4-PlantUML estritamente na notação C4, contendo apenas:  
- Person  
- System_Boundary  
- Component  
- Rel  
- ComponentDb / ComponentDb_Ext  
- ContainerQueue / ContainerQueue_Ext  

Após o diagrama, descreva cada componente: responsabilidades, dependências, dados tratados e quais RNFs o impactam.  

**4) Integrações Externas e Contratos em Alto Nível**  

Liste sistemas externos e descreva protocolos, padrões de integração, formatos, versionamento e objetivos.  

**5) Decisões Arquiteturais Relevantes**  

Registre decisões e trade-offs no formato ADR curto:  
- Contexto → Decisão → Consequências  

**6) Critérios de Aceite e Próximos Passos**  

- Critérios mínimos de aceite.  
- Próximos passos para evolução da arquitetura e implementação.  

---

📌 **Considere também as preferências de stack**:  
{stacks}  

📌 **Texto base para análise**:  
{text}  
"""
    return safe_complete(prompt_content, temp=0.25, max_tokens=4000)


# =============================
# Orquestrador simplificado (só progressbar)
# =============================
def generate_full_doc(text, stacks=""):
    progress = st.progress(0)

    if len(text) < 10000:
        partes = list(split_text(text, max_chars=4000))

        requisitos = []
        for idx, parte in enumerate(partes, 1):
            req_chunk = generate_requisitos_doc(parte)
            requisitos.append(req_chunk)
            progress.progress(idx / (len(partes) + 2))

        requisitos_doc = "\n".join(requisitos)

        tecnico_doc = generate_technical_doc("\n".join(partes), stacks)
        progress.progress((len(partes) + 1) / (len(partes) + 2))

        progress.progress(1.0)

        return f"## Requisitos e Regras de Negócio\n\n{requisitos_doc}\n\n{tecnico_doc}"

    else:
        st.info("🔍 Documento grande detectado. Usando busca semântica (RAG) para otimizar a geração.")

        index = build_index(text)
        progress.progress(0.2)

        requisitos_context = retrieve_relevant_chunks(index, "Requisitos e regras de negócio do sistema")
        progress.progress(0.4)

        tecnico_context = retrieve_relevant_chunks(index, "Aspectos técnicos, arquitetura e integrações")
        progress.progress(0.6)

        requisitos_doc = generate_requisitos_doc(requisitos_context)
        progress.progress(0.75)

        tecnico_doc = generate_technical_doc(tecnico_context, stacks)
        progress.progress(0.95)

        progress.progress(1.0)

        return f"## Requisitos e Regras de Negócio\n\n{requisitos_doc}\n\n{tecnico_doc}"

# =============================
# App Streamlit (UI)
# =============================
st.set_page_config(page_title="Gerador de Documento Técnico", layout="wide")
st.title("📄 Gerador de Documento Técnico Consolidado")

uploaded_file = st.file_uploader("📥 Envie um arquivo PDF ou DOCX (opcional):", type=["pdf", "docx"])
texto_digitado = st.text_area("✍️ Ou cole aqui a história de usuário/caso de uso:", height=220)

st.markdown("### ⚙️ Escolha opcional de Stacks Técnicas")
col1, col2 = st.columns(2)

with col1:
    linguagem = st.selectbox("Linguagem de Programação", ["Não especificar", "Python", "Java", "Node.js", "Go", "C#"])
    banco = st.selectbox("Banco de Dados", ["Não especificar", "PostgreSQL", "MySQL", "MongoDB", "Oracle", "SQL Server"])
    mensageria = st.selectbox("Mensageria/Fila", ["Não especificar", "RabbitMQ", "Kafka", "SQS", "Google Pub/Sub"])

with col2:
    nuvem = st.selectbox("Nuvem", ["Não especificar", "AWS", "Azure", "GCP", "On-premise"])
    arquitetura = st.selectbox("Estilo Arquitetural", ["Não especificar", "Monólito", "Microsserviços", "Serverless", "Event-driven"])

if st.button("🚀 Gerar Documento"):
    texto_base = None
    if uploaded_file:
        texto_base = extract_text_from_file(uploaded_file)
    elif texto_digitado.strip():
        texto_base = texto_digitado.strip()

    stacks = f"""
    Linguagem: {linguagem}
    Banco de Dados: {banco}
    Mensageria: {mensageria}
    Nuvem: {nuvem}
    Estilo Arquitetural: {arquitetura}
    """

    if texto_base:
        with st.spinner("📝 Gerando documento consolidado..."):
            documento = generate_full_doc(texto_base, stacks)
            st.session_state["documento"] = documento
    else:
        st.warning("Forneça um texto ou arquivo para gerar o documento.")

# --- Visualização do documento ---
if "documento" in st.session_state:
    st.subheader("📑 Documento Consolidado")
    st.markdown(st.session_state["documento"])

    st.divider()
    st.markdown("### 🔎 Diagrama Arquitetural")
    plantuml_code = extract_plantuml_code(st.session_state["documento"])
    if plantuml_code:
        svg_content = plantuml_get_diagram(plantuml_code, "svg")
        if svg_content:
            st.components.v1.html(svg_content, height=520, scrolling=True)
    else:
        st.info("Nenhum diagrama detectado no documento.")

    st.divider()
    st.markdown("### 💾 Exportar Documento")
    gerar_download_markdown(st.session_state["documento"])
    gerar_download_pdf(st.session_state["documento"])


